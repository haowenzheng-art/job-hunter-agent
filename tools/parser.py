# tools/parser.py
"""
文档解析器 - 支持 PDF 和 Word 文档
"""
from abc import ABC, abstractmethod
from typing import Optional, List, Dict
from pathlib import Path
import re
from loguru import logger


class BaseParser(ABC):
    """
    文档解析器基类
    """

    def __init__(self):
        self.logger = logger.bind(component="parser")

    @abstractmethod
    def parse(self, file_path: str) -> str:
        """
        解析文档，返回纯文本

        Args:
            file_path: 文件路径

        Returns:
            解析后的文本
        """
        pass

    @abstractmethod
    def is_supported(self, file_path: str) -> bool:
        """
        检查是否支持该文件类型

        Args:
            file_path: 文件路径

        Returns:
            是否支持
        """
        pass

    def clean_text(self, text: str) -> str:
        """
        清理文本（移除多余空格、换行等）

        Args:
            text: 原始文本

        Returns:
            清理后的文本
        """
        # 移除多余空白
        text = re.sub(r'\n\s*\n', '\n\n', text)  # 多个空行变为两个
        text = re.sub(r'[ \t]+', ' ', text)  # 多个空格变为一个

        # 移除页眉页脚（常见模式）
        text = re.sub(r'第\s*\d+\s*页', '', text)
        text = re.sub(r'Page\s*\d+', '', text)

        # 移除特殊字符
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # 标准化换行
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        return text.strip()

    def extract_sections(self, text: str, patterns: Dict[str, str]) -> Dict[str, str]:
        """
        提取文档中的特定章节

        Args:
            text: 文档文本
            patterns: 章节正则模式字典 {章节名: 正则模式}

        Returns:
            提取的章节内容
        """
        sections = {}

        for section_name, pattern in patterns.items():
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                sections[section_name] = match.group(1).strip()

        return sections

    def extract_emails(self, text: str) -> List[str]:
        """提取邮箱地址"""
        return re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)

    def extract_phones(self, text: str) -> List[str]:
        """提取电话号码"""
        patterns = [
            r'1[3-9]\d{9}',  # 中国手机号
            r'\d{3,4}-\d{7,8}',  # 座机
            r'\d{11}',  # 纯数字
        ]
        phones = []
        for pattern in patterns:
            phones.extend(re.findall(pattern, text))
        return list(set(phones))


class PDFParser(BaseParser):
    """
    PDF 文档解析器
    """

    def __init__(self):
        super().__init__()
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖"""
        try:
            from pypdf import PdfReader
            self.PdfReader = PdfReader
        except ImportError:
            raise ImportError(
                "请安装 pypdf: pip install pypdf"
            )

    def is_supported(self, file_path: str) -> bool:
        """检查是否支持该文件类型"""
        return Path(file_path).suffix.lower() in ['.pdf']

    def parse(self, file_path: str) -> str:
        """
        解析 PDF 文档

        Args:
            file_path: PDF 文件路径

        Returns:
            解析后的文本
        """
        if not self.is_supported(file_path):
            raise ValueError(f"不支持的文件类型: {file_path}")

        if not Path(file_path).exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            with open(file_path, 'rb') as file:
                reader = self.PdfReader(file)

                text_parts = []
                total_pages = len(reader.pages)

                self.logger.info(f"开始解析 PDF: {file_path}, 共 {total_pages} 页")

                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                text = '\n'.join(text_parts)
                cleaned = self.clean_text(text)

                self.logger.info(f"PDF 解析完成，共 {len(cleaned)} 字符")
                return cleaned

        except Exception as e:
            self.logger.error(f"PDF 解析失败: {e}")
            raise

    def parse_with_metadata(self, file_path: str) -> Dict[str, any]:
        """
        解析 PDF 并返回元数据

        Args:
            file_path: PDF 文件路径

        Returns:
            包含文本和元数据的字典
        """
        text = self.parse(file_path)

        with open(file_path, 'rb') as file:
            reader = self.PdfReader(file)
            metadata = reader.metadata or {}

        return {
            "text": text,
            "pages": len(reader.pages),
            "title": metadata.get('/Title', ''),
            "author": metadata.get('/Author', ''),
            "subject": metadata.get('/Subject', ''),
            "creator": metadata.get('/Creator', ''),
            "producer": metadata.get('/Producer', '')
        }


class WordParser(BaseParser):
    """
    Word 文档解析器
    """

    def __init__(self):
        super().__init__()
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖"""
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ImportError(
                "请安装 python-docx: pip install python-docx"
            )

    def is_supported(self, file_path: str) -> bool:
        """检查是否支持该文件类型"""
        return Path(file_path).suffix.lower() in ['.docx', '.doc']

    def parse(self, file_path: str) -> bool:
        """
        解析 Word 文档

        Args:
            file_path: Word 文件路径

        Returns:
            解析后的文本
        """
        if not self.is_supported(file_path):
            raise ValueError(f"不支持的文件类型: {file_path}")

        if not Path(file_path).exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # .doc 文件需要额外处理
        if Path(file_path).suffix.lower() == '.doc':
            return self._parse_doc(file_path)

        try:
            doc = self.docx.Document(file_path)

            self.logger.info(f"开始解析 Word 文档: {file_path}")

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            text = '\n'.join(text_parts)
            cleaned = self.clean_text(text)

            self.logger.info(f"Word 文档解析完成，共 {len(cleaned)} 字符")
            return cleaned

        except Exception as e:
            self.logger.error(f"Word 文档解析失败: {e}")
            raise

    def _parse_doc(self, file_path: str) -> str:
        """
        解析 .doc 文件（旧格式）

        Args:
            file_path: .doc 文件路径

        Returns:
            解析后的文本
        """
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            return self.clean_text(text)
        except ImportError:
            raise ImportError(
                "请安装 docx2txt: pip install docx2txt"
            )

    def parse_with_metadata(self, file_path: str) -> Dict[str, any]:
        """
        解析 Word 文档并返回元数据

        Args:
            file_path: Word 文件路径

        Returns:
            包含文本和元数据的字典
        """
        text = self.parse(file_path)

        if Path(file_path).suffix.lower() == '.doc':
            return {
                "text": text,
                "title": "",
                "author": ""
            }

        doc = self.docx.Document(file_path)

        return {
            "text": text,
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
            "created": doc.core_properties.created,
            "modified": doc.core_properties.modified
        }


class TextParser(BaseParser):
    """
    纯文本解析器
    """

    def is_supported(self, file_path: str) -> bool:
        """检查是否支持该文件类型"""
        return Path(file_path).suffix.lower() in ['.txt', '.md', '.rst']

    def parse(self, file_path: str) -> str:
        """
        解析文本文件

        Args:
            file_path: 文本文件路径

        Returns:
            解析后的文本
        """
        if not self.is_supported(file_path):
            raise ValueError(f"不支持的文件类型: {file_path}")

        if not Path(file_path).exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()

            cleaned = self.clean_text(text)
            self.logger.info(f"文本文件解析完成，共 {len(cleaned)} 字符")
            return cleaned

        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'gb2312', 'latin1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    return self.clean_text(text)
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"无法解析文件编码: {file_path}")


class DocumentParserFactory:
    """
    文档解析器工厂
    自动选择合适的解析器
    """

    _parsers = {
        '.pdf': PDFParser,
        '.docx': WordParser,
        '.doc': WordParser,
        '.txt': TextParser,
        '.md': TextParser,
        '.rst': TextParser
    }

    @classmethod
    def get_parser(cls, file_path: str) -> BaseParser:
        """
        根据文件类型获取解析器

        Args:
            file_path: 文件路径

        Returns:
            解析器实例

        Raises:
            ValueError: 不支持的文件类型
        """
        suffix = Path(file_path).suffix.lower()

        if suffix not in cls._parsers:
            raise ValueError(f"不支持的文件类型: {suffix}")

        return cls._parsers[suffix]()

    @classmethod
    def parse(cls, file_path: str) -> str:
        """
        自动解析文档

        Args:
            file_path: 文件路径

        Returns:
            解析后的文本
        """
        parser = cls.get_parser(file_path)
        return parser.parse(file_path)

    @classmethod
    def parse_with_metadata(cls, file_path: str) -> Dict[str, any]:
        """
        自动解析文档并返回元数据

        Args:
            file_path: 文件路径

        Returns:
            包含文本和元数据的字典
        """
        parser = cls.get_parser(file_path)
        if hasattr(parser, 'parse_with_metadata'):
            return parser.parse_with_metadata(file_path)
        else:
            return {
                "text": parser.parse(file_path)
            }

    @classmethod
    def register_parser(cls, suffix: str, parser_class: type):
        """
        注册新的解析器

        Args:
            suffix: 文件后缀
            parser_class: 解析器类
        """
        cls._parsers[suffix] = parser_class